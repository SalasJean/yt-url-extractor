from __future__ import annotations

from pathlib import Path
from unittest.mock import MagicMock, patch

from src.extractor import batch_search, read_document, search_youtube
from src.models import SearchResult


# ------------------------------------------------------------------
# read_document
# ------------------------------------------------------------------
def test_read_txt(tmp_path: Path) -> None:
    file = tmp_path / "canciones.txt"
    content = " Bohemian Rhapsody \n\n Stairway to Heaven \n  \nHotel California\n"
    file.write_text(content, encoding="utf-8")
    lines = read_document(file)
    assert lines == ["Bohemian Rhapsody", "Stairway to Heaven", "Hotel California"]


def test_read_docx(tmp_path: Path) -> None:
    file = tmp_path / "canciones.docx"
    from docx import Document

    doc = Document()
    doc.add_paragraph("Imagine")
    doc.add_paragraph("  Yesterday  ")
    doc.add_paragraph("")
    doc.add_paragraph("Hey Jude")
    doc.save(str(file))

    lines = read_document(file)
    assert lines == ["Imagine", "Yesterday", "Hey Jude"]


def test_read_empty_txt(tmp_path: Path) -> None:
    file = tmp_path / "empty.txt"
    file.write_text("   \n\n  \n", encoding="utf-8")
    assert read_document(file) == []


# ------------------------------------------------------------------
# search_youtube
# ------------------------------------------------------------------
def _make_mock_ydl(entries: list[dict] | None = None) -> MagicMock:
    """Build a YoutubeDL mock whose extract_info returns the given entries."""
    instance = MagicMock()
    instance.extract_info.return_value = {"entries": entries or []}
    return instance


@patch("src.extractor.YoutubeDL")
def test_search_youtube_success(mock_ydl_class: MagicMock) -> None:
    instance = MagicMock()
    instance.extract_info.return_value = {
        "entries": [{"webpage_url": "https://youtube.com/watch?v=ABC123"}],
    }
    mock_ydl_class.return_value.__enter__.return_value = instance

    result = search_youtube("test song")
    assert result.found is True
    assert result.url == "https://youtube.com/watch?v=ABC123"
    assert result.error is None


@patch("src.extractor.YoutubeDL")
def test_search_youtube_no_results(mock_ydl_class: MagicMock) -> None:
    instance = MagicMock()
    instance.extract_info.return_value = {"entries": []}
    mock_ydl_class.return_value.__enter__.return_value = instance

    result = search_youtube("nonexistent")
    assert result.found is False
    assert result.url is None
    assert "No se encontraron resultados" in (result.error or "")


@patch("src.extractor.YoutubeDL")
def test_search_youtube_empty_response(mock_ydl_class: MagicMock) -> None:
    instance = MagicMock()
    instance.extract_info.return_value = {}
    mock_ydl_class.return_value.__enter__.return_value = instance

    result = search_youtube("anything")
    assert result.found is False
    assert result.url is None
    assert "No se encontraron resultados" in (result.error or "")


@patch("src.extractor.YoutubeDL")
def test_search_youtube_network_error(mock_ydl_class: MagicMock) -> None:
    instance = MagicMock()
    instance.extract_info.side_effect = ConnectionError("timeout")
    mock_ydl_class.return_value.__enter__.return_value = instance

    result = search_youtube("fail")
    assert result.found is False
    assert result.url is None
    assert result.error == "timeout"


# ------------------------------------------------------------------
# batch_search
# ------------------------------------------------------------------
@patch("src.extractor.YoutubeDL")
def test_batch_search_callback(mock_ydl_class: MagicMock) -> None:
    instance = MagicMock()
    instance.extract_info.return_value = {
        "entries": [{"webpage_url": "https://youtube.com/watch?v=MOCK"}],
    }
    mock_ydl_class.return_value.__enter__.return_value = instance

    queries = ["song a", "song b", "song c"]
    calls: list[tuple[int, int, SearchResult]] = []

    def cb(current: int, total: int, result: SearchResult) -> None:
        calls.append((current, total, result))

    results = batch_search(queries, progress_callback=cb)

    assert len(results) == 3
    assert [r.found for r in results] == [True, True, True]
    assert len(calls) == 3
    assert calls[0] == (1, 3, results[0])
    assert calls[2] == (3, 3, results[2])
